from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Windows部署和更新操作手册.md"
OUTPUT = ROOT / "Windows部署和更新操作手册.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 40, 55)
MUTED = RGBColor(92, 102, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "DADCE0"


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def add_run(paragraph, text: str, bold=False, color=None, size=None, font="Calibri"):
    run = paragraph.add_run(text)
    set_east_asia_font(run, "Microsoft YaHei" if font == "Calibri" else font)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "杭州特耐时 DXF 智能用料系统 Windows 操作手册", color=MUTED, size=9)
    return doc


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=0, after=3, line=1.15)
    add_run(title, "Windows 部署和更新操作手册", bold=True, color=RGBColor(0, 0, 0), size=24)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12, line=1.25)
    add_run(subtitle, "杭州特耐时 DXF 智能用料系统｜后端服务、小程序配置、数据备份与程序更新", color=MUTED, size=11)

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    widths = [2100, 7260]
    rows = [
        ("适用对象", "在 Windows 电脑上运行系统的使用者。"),
        ("核心原则", "更新程序只更新代码，不覆盖 data、.env、.venv 或 backups。"),
        ("主要入口", "后台：http://127.0.0.1:8000/admin；健康检查：http://127.0.0.1:8000/health"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, LIGHT_BLUE if idx == 0 else "FFFFFF")
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        add_run(row.cells[0].paragraphs[0], label, bold=True, color=DARK_BLUE)
        add_run(row.cells[1].paragraphs[0], value)

    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.1)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.12)
        set_shading(p, LIGHT_GRAY)
        add_run(p, line or " ", font="Consolas", size=9, color=RGBColor(20, 30, 45))
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4, line=1.25)
    add_run(p, text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=4, line=1.25)
    add_run(p, text)


def add_paragraph_text(doc: Document, text: str) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.25)
    add_run(p, text.strip())


def build_doc() -> None:
    doc = setup_document()
    add_title_block(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    skip_first_title = True
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            if skip_first_title:
                skip_first_title = False
                continue
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        bullet_match = re.match(r"^[-*]\s+(.+)$", line)
        if bullet_match:
            add_bullet(doc, bullet_match.group(1).strip())
            continue
        number_match = re.match(r"^\d+\.\s+(.+)$", line)
        if number_match:
            add_number(doc, number_match.group(1).strip())
            continue
        add_paragraph_text(doc, line)

    if code_lines:
        add_code_block(doc, code_lines)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
